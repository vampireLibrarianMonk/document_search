"""Text extraction from PDFs, Word docs, and plain text files.

Handles three types of PDF pages:
  - Text pages: extracted directly with pypdf (free, instant)
  - Image-only pages: sent to vision LLM via Bedrock Converse API (~$0.002/page)
  - Mixed pages: text extracted AND image sent to vision, results merged

Uses the Bedrock Converse API so any supported model works (Claude, Nova, Llama, etc).
"""

from __future__ import annotations

import logging
import os
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

logger = logging.getLogger(__name__)

# If a page has fewer chars than this, we treat it as "no usable text"
MIN_PAGE_TEXT = 20

# Lazy-init so the module imports even without AWS creds
_bedrock = None


def _get_bedrock():
    global _bedrock
    if _bedrock is None:
        import boto3

        _bedrock = boto3.client(
            "bedrock-runtime",
            region_name=os.getenv("AWS_REGION", "us-east-1"),
        )
    return _bedrock


def _page_has_images(page) -> bool:
    """Check if a PDF page contains embedded images."""
    try:
        return len(page.images) > 0
    except Exception:
        try:
            resources = page.get("/Resources", {})
            xobjects = resources.get("/XObject", {})
            return any(xobjects[key].get("/Subtype") == "/Image" for key in xobjects)
        except Exception:
            return False


def _extract_page_image(page, page_num: int, path: str) -> str:
    """Render a PDF page to an image and send it to a vision LLM for OCR."""
    try:
        from io import BytesIO

        from pdf2image import convert_from_path

        # Render just this one page to JPEG (150 DPI keeps it small and cheap)
        images = convert_from_path(
            path,
            first_page=page_num,
            last_page=page_num,
            dpi=150,
            fmt="jpeg",
        )
        if not images:
            return ""

        buf = BytesIO()
        images[0].save(buf, format="JPEG", quality=80)

        model_id = os.getenv("BEDROCK_VISION_MODEL_ID", "mistral.ministral-3-3b-instruct")
        resp = _get_bedrock().converse(
            modelId=model_id,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "image": {
                                "format": "jpeg",
                                "source": {"bytes": buf.getvalue()},
                            },
                        },
                        {
                            "text": "Extract ALL text from this document page, including: company names in logos or headers, watermarks, stamps, handwritten notes, footer text, and any text embedded in images or graphics. Read every piece of visible text on the page. Return only the text content, no commentary.",
                        },
                    ],
                },
            ],
            inferenceConfig={"maxTokens": 4096},
        )
        text = resp["output"]["message"]["content"][0]["text"]

        # Track usage if enabled
        if os.getenv("TRACK_USAGE", "true").lower() == "true":
            usage = resp.get("usage", {})
            from .db import get_conn
            from .pricing import estimate_cost

            try:
                cost = estimate_cost(
                    model_id,
                    usage.get("inputTokens", 0),
                    usage.get("outputTokens", 0),
                    os.getenv("AWS_REGION", "us-east-1"),
                )
                conn = get_conn()
                with conn.cursor() as cur:
                    cur.execute(
                        """INSERT INTO token_usage
                           (model_id, operation, input_tokens, output_tokens, estimated_cost_usd)
                           VALUES (%s, %s, %s, %s, %s)""",
                        (
                            model_id,
                            "vision",
                            usage.get("inputTokens", 0),
                            usage.get("outputTokens", 0),
                            cost,
                        ),
                    )
                conn.close()
            except Exception:
                pass  # nosec B110  # nosec B110 - non-fatal, logged elsewhere

        logger.info(
            "Vision OCR extracted %d chars from page %d of %s",
            len(text),
            page_num,
            Path(path).name,
        )
        return text
    except Exception as e:
        logger.warning("Vision OCR failed for page %d of %s: %s", page_num, Path(path).name, e)
        return ""


def extract_text(path: str) -> tuple[str, list[str]]:
    """Pull plain text out of a PDF, DOCX, image, or text file.

    Returns (extracted_text, processing_log) where processing_log is a list
    of human-readable messages describing what happened during extraction.

    For PDFs, works per-page:
      - Text-only page: use extracted text
      - Image-only page (no text): send to vision LLM
      - Mixed page (text + images): extract text AND send to vision, merge both

    For DOCX/DOC, walks the XML tree in document order:
      - Text runs are extracted directly
      - Inline images are sent to vision LLM at their exact position

    For standalone images (.jpg, .png, .tiff):
      - Sent directly to vision LLM for OCR
    """
    ext = Path(path).suffix.lower()
    log: list[str] = []

    if ext == ".pdf":
        reader = PdfReader(path)
        log.append(f"PDF with {len(reader.pages)} pages")
        pages_text = []
        for i, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            has_text = len(text) >= MIN_PAGE_TEXT
            has_images = _page_has_images(page)

            if has_text and not has_images:
                pages_text.append(text)
                log.append(f"Page {i}: extracted text ({len(text)} chars)")
            elif has_text and has_images:
                ocr_text = _extract_page_image(page, i, path)
                if ocr_text and ocr_text.strip() != text.strip():
                    pages_text.append(f"{text}\n\n[Image content]\n{ocr_text}")
                    log.append(f"Page {i}: text + vision OCR ({len(text)} + {len(ocr_text)} chars)")
                else:
                    pages_text.append(text)
                    log.append(f"Page {i}: text only, image had no new content ({len(text)} chars)")
            else:
                ocr_text = _extract_page_image(page, i, path)
                if ocr_text:
                    pages_text.append(ocr_text)
                    log.append(f"Page {i}: scanned, vision OCR ({len(ocr_text)} chars)")
                else:
                    log.append(f"Page {i}: empty, no text or image content")
        result = "\n".join(pages_text).strip().replace("\x00", "")
        return result, log

    if ext in (".docx", ".doc"):
        result, doc_log = _extract_docx_with_images(path)
        log.extend(doc_log)
        return result.replace("\x00", ""), log

    if ext in (".jpg", ".jpeg", ".png", ".tiff", ".tif"):
        log.append("Standalone image, sending to vision LLM")
        ocr_text = _extract_standalone_image(path)
        if ocr_text:
            log.append(f"Vision OCR extracted {len(ocr_text)} chars")
        else:
            log.append("Vision OCR returned no text")
        return (ocr_text or "").replace("\x00", ""), log

    # .txt, .md, or anything else
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read().strip()
    log.append(f"Plain text file ({len(text)} chars)")
    return text, log


def _extract_standalone_image(path: str) -> str:
    """Send a standalone image file to the vision LLM for OCR."""
    try:
        with open(path, "rb") as f:
            img_bytes = f.read()

        # Determine format from extension
        ext = Path(path).suffix.lower()
        fmt_map = {".jpg": "jpeg", ".jpeg": "jpeg", ".png": "png", ".tiff": "tiff", ".tif": "tiff"}
        fmt = fmt_map.get(ext, "jpeg")

        model_id = os.getenv("BEDROCK_VISION_MODEL_ID", "mistral.ministral-3-3b-instruct")
        resp = _get_bedrock().converse(
            modelId=model_id,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"image": {"format": fmt, "source": {"bytes": img_bytes}}},
                        {"text": "Extract ALL text from this document image, including: company names in logos or headers, watermarks, stamps, handwritten notes, footer text, and any text embedded in images or graphics. Read every piece of visible text on the page. Return only the text content, no commentary."},
                    ],
                },
            ],
            inferenceConfig={"maxTokens": 4096},
        )
        text = resp["output"]["message"]["content"][0]["text"]

        # Track usage
        if os.getenv("TRACK_USAGE", "true").lower() == "true":
            usage = resp.get("usage", {})
            from .db import get_conn
            from .pricing import estimate_cost

            try:
                cost = estimate_cost(
                    model_id,
                    usage.get("inputTokens", 0),
                    usage.get("outputTokens", 0),
                    os.getenv("AWS_REGION", "us-east-1"),
                )
                conn = get_conn()
                with conn.cursor() as cur:
                    cur.execute(
                        """INSERT INTO token_usage
                           (model_id, operation, input_tokens, output_tokens, estimated_cost_usd)
                           VALUES (%s, %s, %s, %s, %s)""",
                        (
                            model_id,
                            "vision",
                            usage.get("inputTokens", 0),
                            usage.get("outputTokens", 0),
                            cost,
                        ),
                    )
                conn.close()
            except Exception:
                pass  # nosec B110

        return text
    except Exception as e:
        logger.warning("Standalone image OCR failed for %s: %s", path, e)
        return ""


def _extract_docx_with_images(path: str) -> tuple[str, list[str]]:
    """Extract text and images from a DOCX in document order.

    Walks the XML tree so images appear at their exact position in the text.
    Falls back to text-only extraction if image processing fails.
    """
    from docx.oxml.ns import qn

    log: list[str] = []
    doc = DocxDocument(path)
    parts: list[str] = []
    image_count = 0

    # Get the document's relationship part for resolving image references
    rels = doc.part.rels

    log.append(f"DOCX with {len(doc.paragraphs)} paragraphs")

    for para in doc.paragraphs:
        para_parts: list[str] = []

        for run in para.runs:
            # Check if this run contains a drawing (image)
            drawings = run.element.findall(f".//{qn('w:drawing')}")
            if drawings:
                for drawing in drawings:
                    # Extract the image relationship ID
                    blips = drawing.findall(f".//{qn('a:blip')}")
                    for blip in blips:
                        embed_id = blip.get(qn("r:embed"))
                        if embed_id and embed_id in rels:
                            try:
                                image_part = rels[embed_id].target_part
                                img_bytes = image_part.blob
                                image_count += 1

                                # Send to vision LLM
                                model_id = os.getenv("BEDROCK_VISION_MODEL_ID", "mistral.ministral-3-3b-instruct")
                                if model_id:
                                    # Determine format
                                    content_type = image_part.content_type or "image/png"
                                    fmt = content_type.split("/")[-1]
                                    if fmt == "jpeg":
                                        fmt = "jpeg"
                                    elif fmt in ("png", "gif", "webp"):
                                        pass
                                    else:
                                        fmt = "png"

                                    resp = _get_bedrock().converse(
                                        modelId=model_id,
                                        messages=[
                                            {
                                                "role": "user",
                                                "content": [
                                                    {"image": {"format": fmt, "source": {"bytes": img_bytes}}},
                                                    {"text": "Extract ALL text from this image, including: company names in logos or headers, watermarks, stamps, handwritten notes, and any text embedded in graphics. Return only the text content, no commentary."},
                                                ],
                                            },
                                        ],
                                        inferenceConfig={"maxTokens": 4096},
                                    )
                                    ocr = resp["output"]["message"]["content"][0]["text"]
                                    para_parts.append(f"\n[Image {image_count}]\n{ocr}")
                                    log.append(f"Image {image_count}: vision OCR ({len(ocr)} chars)")

                                    # Track usage
                                    if os.getenv("TRACK_USAGE", "true").lower() == "true":
                                        usage = resp.get("usage", {})
                                        from .db import get_conn
                                        from .pricing import estimate_cost

                                        try:
                                            cost = estimate_cost(
                                                model_id,
                                                usage.get("inputTokens", 0),
                                                usage.get("outputTokens", 0),
                                                os.getenv("AWS_REGION", "us-east-1"),
                                            )
                                            conn = get_conn()
                                            with conn.cursor() as cur:
                                                cur.execute(
                                                    """INSERT INTO token_usage
                                                       (model_id, operation, input_tokens, output_tokens, estimated_cost_usd)
                                                       VALUES (%s, %s, %s, %s, %s)""",
                                                    (
                                                        model_id,
                                                        "vision",
                                                        usage.get("inputTokens", 0),
                                                        usage.get("outputTokens", 0),
                                                        cost,
                                                    ),
                                                )
                                            conn.close()
                                        except Exception:
                                            pass  # nosec B110
                                else:
                                    log.append(f"Image {image_count}: skipped (no vision model configured)")
                            except Exception as e:
                                log.append(f"Image {image_count}: extraction failed ({e})")
            # Always include the run's text
            if run.text:
                para_parts.append(run.text)

        if para_parts:
            parts.append("".join(para_parts))

    if image_count == 0:
        log.append("No images found in document")
    else:
        log.append(f"Processed {image_count} inline images")

    return "\n".join(parts).strip(), log


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
    """Split text into overlapping windows, breaking on sentence boundaries.

    Each chunk is roughly chunk_size characters with overlap characters
    shared between consecutive chunks so context isn't lost at boundaries.
    """
    if not text:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + chunk_size)
        # Try to break on a sentence boundary instead of mid-word
        if end < len(text):
            for sep in (". ", ".\n", "\n\n", "\n", " "):
                boundary = text.rfind(sep, start + chunk_size // 2, end)
                if boundary != -1:
                    end = boundary + len(sep)
                    break
        chunks.append(text[start:end].strip())
        if end >= len(text):
            break
        start = max(start + 1, end - overlap)
    return [c for c in chunks if c]
